"""
SealAI RAG Ingest (Qdrant)

Used by:
- async job worker (expects: ingest_file(...))
- optional CLI usage (python3 rag_ingest.py <path> --tenant ...)

Goals:
- Stable worker-compatible API: ingest_file(...)
- Payload schema: payload["metadata"].* is SoT (matches metadata.* filters)
- Keep backward-compatible top-level fields (text/source/filename/tenant_id/document_id/visibility)
- Respect env config for collection/vector naming and embedding models
- Deterministic chunk IDs (tenant_id + document_id + chunk_index)
- Safe text loading + bounded chunk sizes
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import logging
import mimetypes
import os
import re
import time
import zipfile
from dataclasses import dataclass
from typing import Any, Dict, Iterable, List, Optional

from qdrant_client import QdrantClient, models

from langchain_community.document_loaders import Docx2txtLoader

try:
    from langchain_qdrant import QdrantVectorStore
except ImportError:
    QdrantVectorStore = None

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    HuggingFaceEmbeddings = None


def get_embedder() -> Any:
    raise NotImplementedError("Use IngestPipeline instead of the legacy helper.")

from app.services.rag.rag_schema import ChunkMetadata, Domain, EngineeringProps, MaterialFamily, SourceType, TempRange
from app.services.rag.document import Document
from app.services.rag.rag_etl_pipeline import (
    LLMDocumentExtraction, LLMOperatingPoint, LLMCondition, LLMLimit,
    Operator, process_document_pipeline, PipelineStatus
)
from app.services.rag.qdrant_state_machine import transition_to_published_bulletproof

log = logging.getLogger(__name__)

# -----------------------------------------------------------------------------
# Config (ENV-driven)
# -----------------------------------------------------------------------------

QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY") or None

# Collection must be consistent across worker + retrieval.
COLLECTION_NAME = os.getenv("QDRANT_COLLECTION", "sealai_knowledge")

# Qdrant vector naming:
# - "" (empty) means: single unnamed vector collection (vectors: {"": {...}})
# - "dense" means: named vector collection (vectors: {"dense": {...}})
QDRANT_VECTOR_NAME = os.getenv("QDRANT_VECTOR_NAME", "")

# Embedding models (override via env)
DENSE_MODEL = (
    os.getenv("RAG_DENSE_MODEL")
    or os.getenv("embedding_model")
    or "BAAI/bge-base-en-v1.5"
).strip()
SPARSE_MODEL = os.getenv("RAG_SPARSE_MODEL", "prithivida/Splade_PP_en_v1")

# Chunking
MAX_CHUNK_CHARS   = int(os.getenv("RAG_MAX_CHUNK_CHARS", "6000"))
RAG_CHUNK_OVERLAP = int(os.getenv("RAG_CHUNK_OVERLAP", "200"))

# Sparse vectors only if the target collection supports it.
ENABLE_SPARSE = os.getenv("RAG_SPARSE_ENABLED", "0").strip().lower() not in ("0", "false", "no")
LEGACY_VECTORSTORE_ENABLED = os.getenv("RAG_INGEST_LEGACY_VECTORSTORE", "0").strip().lower() in (
    "1",
    "true",
    "yes",
    "on",
)
DEFAULT_INGEST_TENANT = (os.getenv("RAG_INGEST_DEFAULT_TENANT") or "sealai").strip() or "sealai"
RAG_SHARED_TENANT_ENABLED = os.getenv("RAG_SHARED_TENANT_ENABLED", "0").strip().lower() in ("1", "true", "yes", "on")
RAG_SHARED_TENANT_ID = (os.getenv("RAG_SHARED_TENANT_ID") or "").strip()
RAG_DYNAMIC_METADATA_LLM_ENABLED = os.getenv("RAG_DYNAMIC_METADATA_LLM_ENABLED", "1").strip().lower() not in (
    "0",
    "false",
    "no",
)
RAG_DYNAMIC_METADATA_LLM_MODEL = (os.getenv("RAG_DYNAMIC_METADATA_LLM_MODEL") or "gpt-4.1-mini").strip()
RAG_DYNAMIC_METADATA_MAX_CHARS = int(os.getenv("RAG_DYNAMIC_METADATA_MAX_CHARS", "12000"))


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------

def _ensure_tenant_allowed(tenant_id: str) -> None:
    if not tenant_id or not str(tenant_id).strip():
        raise ValueError("tenant_id required for ingest")
    if tenant_id == "sealai":
        shared_allowed = bool(RAG_SHARED_TENANT_ENABLED and RAG_SHARED_TENANT_ID == "sealai")
        if not shared_allowed:
            raise ValueError(
                "tenant_id 'sealai' is reserved; set RAG_SHARED_TENANT_ENABLED=1 and "
                "RAG_SHARED_TENANT_ID=sealai to allow shared ingestion"
            )


def _coerce_visibility(v: str | None) -> str:
    v = (v or "public").strip().lower()
    return "private" if v == "private" else "public"


def _safe_read_text_file(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="latin-1", errors="replace") as f:
            return f.read()
    except Exception as e:
        print(f"[WARN] Could not read text file {path}: {e}")
        return ""


def _extract_docx_text_via_zip(file_path: str) -> str:
    try:
        with zipfile.ZipFile(file_path) as archive:
            raw = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        normalized = (
            raw.replace("</w:p>", "\n")
            .replace("</w:tr>", "\n")
            .replace("</w:tc>", "\t")
        )
        text = re.sub(r"<[^>]+>", "", normalized)
        return html.unescape(text)
    except Exception:
        return ""


def _extract_docx_text_python_docx(file_path: str) -> str:
    try:
        import docx  # type: ignore
    except Exception:
        return ""

    try:
        doc = docx.Document(file_path)
        parts: List[str] = []

        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [(cell.text or "").strip() for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    parts.append(row_text)

        # Optional headers/footers can contain relevant content in technical docs.
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                text = (paragraph.text or "").strip()
                if text:
                    parts.append(text)
            for paragraph in section.footer.paragraphs:
                text = (paragraph.text or "").strip()
                if text:
                    parts.append(text)

        return "\n".join(parts)
    except Exception:
        return ""


def _load_pdf_pages(file_path: str) -> List[tuple[int, str]]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"pypdf_unavailable: {type(exc).__name__}: {exc}") from exc

    try:
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            if reader.is_encrypted:
                try:
                    decrypted = reader.decrypt("")  # type: ignore[call-arg]
                except Exception:
                    decrypted = 0
                if not decrypted:
                    raise ValueError(f"encrypted_pdf_not_supported: {file_path}")

            pages: List[tuple[int, str]] = []
            has_text = False
            for idx, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                except Exception as exc:
                    print(f"[WARN] Could not extract PDF page {idx + 1} from {file_path}: {exc}")
                    page_text = ""
                if page_text.strip():
                    has_text = True
                pages.append((idx + 1, page_text))

            if not has_text:
                print(f"[WARN] No text extracted from PDF {file_path}")
                return []
            return pages
    except ValueError:
        raise
    except Exception as exc:
        print(f"[WARN] Could not load PDF {file_path}: {exc}")
        return []


def _load_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".docx":
            text = _extract_docx_text_python_docx(file_path)
            if text.strip():
                return text
            try:
                loader = Docx2txtLoader(file_path)
                docs = loader.load()
                text = "\n\n".join([d.page_content for d in docs if getattr(d, "page_content", None)])
                if text.strip():
                    return text
            except Exception:
                pass
            fallback = _extract_docx_text_via_zip(file_path)
            if fallback.strip():
                return fallback
            log.warning("docx_text_empty_after_extraction", extra={"file_path": file_path})
            return ""
        return _safe_read_text_file(file_path)
    except Exception as e:
        print(f"[WARN] Could not load {file_path}: {e}")
        return ""


def _load_pages(file_path: str) -> List[tuple[Optional[int], str]]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _load_pdf_pages(file_path)
    return [(None, _load_text(file_path))]


def load_document(file_path: str) -> List[Any]:
    from langchain_core.documents import Document

    docs: List[Any] = []
    for page_number, text in _load_pages(file_path):
        if not str(text or "").strip():
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "page": page_number,
                },
            )
        )
    return docs


def _chunk_pages(pages: List[tuple[Optional[int], str]], filename: str = "") -> List[tuple[str, Optional[int]]]:
    chunks: List[tuple[str, Optional[int]]] = []
    for page_number, text in pages:
        if not text.strip():
            continue
        for chunk in _semantic_paragraph_chunks(text):
            if chunk:
                # Platinum Blueprint v4.1: Prepend context to every chunk
                prefix = f"[Document: {filename}] " if filename else ""
                chunks.append((f"{prefix}{chunk}", page_number))
    return chunks


def _semantic_paragraph_chunks(text: str) -> List[str]:
    """
    Splits text into chunks with overlap. 
    Prioritizes paragraph breaks but ensures sliding window overlap to keep context.
    """
    if not text.strip():
        return []
        
    # Joint based sliding window for technical context retention
    chunks: List[str] = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = start + MAX_CHUNK_CHARS
        chunk = text[start:end]
        
        # Try to snap to the last paragraph break within the chunk to keep it clean
        if end < text_len:
            last_para = chunk.rfind("\n\n")
            if last_para > MAX_CHUNK_CHARS // 2:
                chunk = text[start : start + last_para]
                end = start + last_para
        
        chunks.append(chunk.strip())
        start = end - RAG_CHUNK_OVERLAP
        if start >= text_len or end >= text_len:
            break
            
    return [c for c in chunks if c]


DOMAIN_MAPPING = {
    "Norm": "standard",
    "Product": "product",
    "Standard": "standard",
    "Technical": "material",
    "Datasheet": "material",
}

_MATERIAL_CODE_PATTERN = re.compile(r"\b([A-Z]{2,6}[-_/]?\d{2,4})\b", re.IGNORECASE)
_MATERIAL_FAMILY_PATTERN = re.compile(
    r"\b(NBR|HNBR|FKM|FFKM|EPDM|PTFE|VMQ|FVMQ|PU|PUR|TPU|PEEK)\b",
    re.IGNORECASE,
)
_SHORE_PATTERN = re.compile(r"\b(?:shore(?:\s*[ad])?|sh)\s*[:=]?\s*(\d{2,3})\b", re.IGNORECASE)
_TEMP_RANGE_PATTERN = re.compile(
    r"\b(-?\d{1,3}(?:[.,]\d+)?)\s*(?:\.\.|to|bis|[-–—])\s*(-?\d{1,3}(?:[.,]\d+)?)\s*°?\s*c?\b",
    re.IGNORECASE,
)
_UNKNOWN_MATERIAL_CODE = "UNKNOWN"
_DEFAULT_SHORE_HARDNESS = 70
_DEFAULT_TEMP_RANGE = {"min_c": -40.0, "max_c": 120.0}

_DYNAMIC_NUMERIC_PATTERNS: Dict[str, re.Pattern[str]] = {
    "density_kg_m3": re.compile(
        r"(?:dichte|density)\s*(?:von|:|=)?\s*(\d{2,6}(?:[.,]\d+)?)\s*kg\s*/?\s*m(?:\^?3|³)",
        re.IGNORECASE,
    ),
    "melting_point_c": re.compile(
        r"(?:schmelzpunkt|melting point)\s*(?:von|:|=)?\s*(-?\d{1,4}(?:[.,]\d+)?)\s*°\s*c",
        re.IGNORECASE,
    ),
    "glass_transition_temp_c": re.compile(
        r"(?:glasübergangstemperatur|glass transition(?: temperature)?)\s*(?:von|:|=)?\s*(-?\d{1,4}(?:[.,]\d+)?)\s*°\s*c",
        re.IGNORECASE,
    ),
    "thermal_conductivity_w_mk": re.compile(
        r"(?:wärmeleitfähigkeit|thermal conductivity)\s*(?:von|:|=)?\s*(-?\d{1,4}(?:[.,]\d+)?)\s*w\s*/?\s*\(?(?:m[·*]k|m/k)\)?",
        re.IGNORECASE,
    ),
    "elastic_modulus_gpa": re.compile(
        r"(?:elastizitätsmodul|young'?s modulus|elastic modulus)\s*(?:von|:|=)?\s*(-?\d{1,4}(?:[.,]\d+)?)\s*gpa",
        re.IGNORECASE,
    ),
    "yield_strength_mpa": re.compile(
        r"(?:streckgrenze|yield strength)\s*(?:von|:|=)?\s*(-?\d{1,4}(?:[.,]\d+)?)\s*mpa",
        re.IGNORECASE,
    ),
    "decomposition_start_c": re.compile(
        r"(?:beginnt bei|decompose(?:s|d)?(?: above| over)?|zersetz\w+\s*(?:bei|ab|über|ueber))\s*(-?\d{1,4}(?:[.,]\d+)?)\s*°\s*c",
        re.IGNORECASE,
    ),
}
_FRICTION_RANGE_PATTERN = re.compile(
    r"(?:reibungskoeffizient|friction coefficient)\s*(?:von|:|=)?\s*(-?\d+(?:[.,]\d+)?)\s*(?:to|bis|[-–—])\s*(-?\d+(?:[.,]\d+)?)",
    re.IGNORECASE,
)

def _parse_domain(domain_str: str | None) -> Domain:
    if not domain_str:
        return Domain.MATERIAL
    s = str(domain_str).strip()
    if not s:
        return Domain.MATERIAL

    # Apply mapping
    mapped = DOMAIN_MAPPING.get(s)
    if mapped:
        s = mapped

    # Enum name
    try:
        upper = s.upper()
        if hasattr(Domain, "__members__") and upper in Domain.__members__:
            return Domain.__members__[upper]  # type: ignore[index]
    except Exception:
        pass

    # Enum value
    try:
        return Domain(s)  # type: ignore[arg-type]
    except Exception:
        return Domain.MATERIAL


def _hash_path(path: str) -> str:
    return hashlib.md5(path.encode("utf-8")).hexdigest()[:12]


def _parse_scalar(value: str) -> object:
    raw = (value or "").strip()
    if not raw:
        return raw
    low = raw.lower()
    if low in {"true", "yes", "ja"}:
        return True
    if low in {"false", "no", "nein"}:
        return False
    if re.fullmatch(r"-?\d+", raw):
        try:
            return int(raw)
        except ValueError:
            return raw
    if re.fullmatch(r"-?\d+[.,]\d+", raw):
        try:
            return float(raw.replace(",", "."))
        except ValueError:
            return raw
    return raw


def _parse_facets_from_tags(tags: Iterable[str] | None) -> dict[str, object]:
    facets = {
        "entity": None,
        "aspects": [],
        "language": None,
        "source_version": None,
        "effective_date": None,
        "material_code": None,
        "source_url": None,
        "shore_hardness": None,
        "temp_range": {"min_c": None, "max_c": None},
        "additional_metadata": {},
    }
    seen_aspects: set[str] = set()
    for raw in tags or []:
        if not raw or ":" not in raw:
            continue
        key, value = raw.split(":", 1)
        key = key.strip().lower()
        value = value.strip()
        if not value:
            continue
        if key in {"entity", "material"}:
            facets["entity"] = value
        elif key in {"aspect", "aspects"}:
            normalized = value.lower()
            if normalized not in seen_aspects:
                seen_aspects.add(normalized)
                facets["aspects"].append(value)
        elif key in {"lang", "language"}:
            facets["language"] = value
        elif key in {"version", "source_version"}:
            facets["source_version"] = value
        elif key in {"effective", "effective_date"}:
            facets["effective_date"] = value
        elif key in {"material_code", "material", "code"}:
            facets["material_code"] = value
        elif key in {"source_url", "url", "source"}:
            facets["source_url"] = value
        elif key in {"shore_hardness", "shore"}:
            try:
                facets["shore_hardness"] = int(value)
            except ValueError:
                pass
        elif key in {"temp_range", "temperature_range"}:
            parsed = _parse_temp_range_text(value)
            if parsed:
                facets["temp_range"] = parsed
        elif key in {"temp_min", "temperature_min"}:
            try:
                facets["temp_range"]["min_c"] = float(value.replace(",", "."))
            except (ValueError, TypeError):
                pass
        elif key in {"temp_max", "temperature_max"}:
            try:
                facets["temp_range"]["max_c"] = float(value.replace(",", "."))
            except (ValueError, TypeError):
                pass
        else:
            additional = facets.get("additional_metadata")
            if isinstance(additional, dict):
                additional[key] = _parse_scalar(value)
    return facets


def _guess_entity_from_filename(filename: str) -> Optional[str]:
    name = os.path.splitext(filename or "")[0].upper()
    for member in MaterialFamily:
        if member.value.upper() in name:
            return member.value
    return None


def _parse_temp_range_text(value: str | None) -> Optional[dict[str, Optional[float]]]:
    text = (value or "").strip()
    if not text:
        return None
    match = _TEMP_RANGE_PATTERN.search(text.lower())
    if not match:
        return None
    try:
        min_c = float(match.group(1).replace(",", "."))
        max_c = float(match.group(2).replace(",", "."))
    except (TypeError, ValueError):
        return None
    return {"min_c": min_c, "max_c": max_c}


def _normalize_material_code(value: Optional[str]) -> str:
    candidate = (value or "").strip().upper()
    if candidate:
        return candidate
    return _UNKNOWN_MATERIAL_CODE


def _normalize_temp_range_payload(
    payload: object,
    *texts: Optional[str],
) -> dict[str, float]:
    candidate: Optional[dict[str, Optional[float]]] = None
    if isinstance(payload, dict):
        min_c_raw = payload.get("min_c")
        max_c_raw = payload.get("max_c")
        try:
            min_c = float(min_c_raw) if min_c_raw is not None else None
            max_c = float(max_c_raw) if max_c_raw is not None else None
            candidate = {"min_c": min_c, "max_c": max_c}
        except (TypeError, ValueError):
            candidate = None
    if not candidate or candidate.get("min_c") is None or candidate.get("max_c") is None:
        for text in texts:
            parsed = _parse_temp_range_text(text)
            if parsed and parsed.get("min_c") is not None and parsed.get("max_c") is not None:
                candidate = parsed
                break
    if not candidate or candidate.get("min_c") is None or candidate.get("max_c") is None:
        candidate = dict(_DEFAULT_TEMP_RANGE)
    lo = float(candidate["min_c"])
    hi = float(candidate["max_c"])
    if lo > hi:
        lo, hi = hi, lo
    return {"min_c": lo, "max_c": hi}


def _extract_material_code(*values: Optional[str]) -> Optional[str]:
    for value in values:
        text = (value or "").strip()
        if not text:
            continue
        match = _MATERIAL_CODE_PATTERN.search(text)
        if match:
            return match.group(1).upper()
        family = _MATERIAL_FAMILY_PATTERN.search(text)
        if family:
            return family.group(1).upper()
    return None


def _extract_shore_hardness(*values: Optional[str]) -> Optional[int]:
    for value in values:
        text = (value or "").strip()
        if not text:
            continue
        match = _SHORE_PATTERN.search(text)
        if match:
            try:
                hardness = int(match.group(1))
                if 0 <= hardness <= 120:
                    return hardness
            except (TypeError, ValueError):
                continue
    return None


def _extract_dynamic_metadata_regex(
    *,
    text: str,
    filename: str,
) -> Dict[str, Any]:
    dynamic: Dict[str, Any] = {}
    haystack = f"{filename}\n{text}"
    for key, pattern in _DYNAMIC_NUMERIC_PATTERNS.items():
        match = pattern.search(haystack)
        if not match:
            continue
        try:
            dynamic[key] = float(match.group(1).replace(",", "."))
        except (TypeError, ValueError):
            continue

    friction_match = _FRICTION_RANGE_PATTERN.search(haystack)
    if friction_match:
        try:
            lo = float(friction_match.group(1).replace(",", "."))
            hi = float(friction_match.group(2).replace(",", "."))
            dynamic["friction_coefficient_min"] = min(lo, hi)
            dynamic["friction_coefficient_max"] = max(lo, hi)
        except (TypeError, ValueError):
            pass

    if re.search(r"\bpolytetrafluoroethylene\b", haystack, re.IGNORECASE):
        dynamic["polymer_name"] = "Polytetrafluoroethylene"
    if re.search(r"\bkyrolon\s+79x\b", haystack, re.IGNORECASE):
        dynamic["trade_name"] = "Kyrolon 79X"
    if re.search(r"\bhydrophob", haystack, re.IGNORECASE):
        dynamic["hydrophobic"] = True
    if re.search(r"\bchemisch inert|chemical inert", haystack, re.IGNORECASE):
        dynamic["chemically_inert"] = True
    return dynamic



def _extract_dynamic_metadata_llm(
    *,
    text: str,
    filename: str,
    seed: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Lightweight LLM-based metadata enrichment for non-PDF documents (DOCX, TXT).
    Extracts generic entity/material metadata in JSON mode.
    Soft-fails (returns seed) if OPENAI_API_KEY is missing or the call fails.
    NOTE: For PDFs, use _extract_platinum_structured_llm instead.
    """
    result: Dict[str, Any] = dict(seed or {})

    if not RAG_DYNAMIC_METADATA_LLM_ENABLED:
        return result

    api_key = (os.getenv("OPENAI_API_KEY") or "").strip()
    if not api_key:
        log.debug("dynamic_metadata_llm_skipped_no_key", extra={"doc_filename": filename})
        return result

    try:
        from openai import OpenAI  # type: ignore
        excerpt = (text or "").strip()[:RAG_DYNAMIC_METADATA_MAX_CHARS]

        client = OpenAI(api_key=api_key)
        system = (
            "Du bist ein Metadaten-Extraktions-Assistent für Dichtungs-Datenblaetter. "
            "Antworte ausschliesslich mit einem validen JSON-Objekt (kein Markdown). "
            "Moegliche Felder: entity (string, Produktname), material_family (string), "
            "hardness_shore (number), trade_name (string), polymer_name (string). "
            "Lasse Felder weg, die du nicht mit Sicherheit bestimmen kannst."
        )
        response = client.chat.completions.create(
            model=RAG_DYNAMIC_METADATA_LLM_MODEL,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": f"Datei: {filename}\nText:\n{excerpt}"},
            ],
            temperature=0,
            max_tokens=512,
        )
        raw = (response.choices[0].message.content or "").strip()
        extracted: Dict[str, Any] = json.loads(raw)
        # Only merge keys not already present from regex/tags
        for k, v in extracted.items():
            result.setdefault(k, v)
        result["dynamic_metadata_extraction"] = "llm"
    except Exception as exc:
        log.warning(
            "dynamic_metadata_llm_failed",
            extra={"doc_filename": filename, "error": str(exc)},
        )

    return result


def _extract_platinum_structured_llm(
    *,
    text: str,
    filename: str,
) -> LLMDocumentExtraction:
    """
    V5.2 PED-Platinum: Uses Vision-LLM to extract structured operating points.
    """
    api_key = (os.getenv("OPENAI_API_KEY") or "").strip()
    if not api_key:
        raise ValueError("OPENAI_API_KEY missing for Platinum Extraction.")

    from openai import OpenAI  # type: ignore
    excerpt = (text or "").strip()
    if len(excerpt) > RAG_DYNAMIC_METADATA_MAX_CHARS:
        excerpt = excerpt[:RAG_DYNAMIC_METADATA_MAX_CHARS]
    
    client = OpenAI(api_key=api_key)
    
    system = (
        "Du bist ein Senior Sealing Engineer. Extrahiere technische Operating Points aus dem Datenblatt. "
        "Nutze das bereitgestellte Schema. Sei präzise bei Units und Evidence-Referenzen (Seite/Abschnitt)."
    )
    
    # Wir nutzen hier Pydantic-based response format für maximale Stabilität (Blueprint v4.1)
    try:
        response = client.beta.chat.completions.parse(
            model=RAG_DYNAMIC_METADATA_LLM_MODEL,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": f"Datei: {filename}\nText:\n{excerpt}"}
            ],
            response_format=LLMDocumentExtraction
        )
        return response.choices[0].message.parsed
    except Exception as exc:
        log.error("platinum_extraction_failed", extra={"doc_filename": filename, "error": str(exc)})
        raise


def _extract_additional_metadata(
    *,
    text: str,
    filename: str,
    tag_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    merged: Dict[str, Any] = dict(tag_metadata or {})
    merged.update(_extract_dynamic_metadata_regex(text=text, filename=filename))
    merged = _extract_dynamic_metadata_llm(text=text, filename=filename, seed=merged)
    merged.setdefault("dynamic_metadata_extraction", "regex")
    return merged


@dataclass
class IngestStats:
    chunks: int
    elapsed_ms: int
    file_size: Optional[int] = None


# -----------------------------------------------------------------------------
# Pipeline
# -----------------------------------------------------------------------------

class IngestPipeline:
    def __init__(self) -> None:
        self.client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY)
        self.dense_model = DENSE_MODEL
        self.sparse_model = SPARSE_MODEL
        self.sparse_enabled = ENABLE_SPARSE
        self._dense_embedder = None
        self._sparse_embedder = None
        self._sparse_capability_checked = False

    def _ensure_sparse_compatibility(self) -> None:
        if self._sparse_capability_checked or not self.sparse_enabled:
            self._sparse_capability_checked = True
            return
        self._sparse_capability_checked = True
        try:
            info = self.client.get_collection(COLLECTION_NAME)
            config = getattr(info, "config", None)
            params = getattr(config, "params", None) if config is not None else None
            sparse_vectors = getattr(params, "sparse_vectors", None) if params is not None else None
            has_sparse = False
            if isinstance(sparse_vectors, dict):
                has_sparse = "sparse" in sparse_vectors
            elif sparse_vectors is not None:
                has_sparse = bool(getattr(sparse_vectors, "sparse", None))
            if not has_sparse:
                print(
                    f"[WARN] Sparse ingest disabled for {COLLECTION_NAME}: "
                    "collection has no 'sparse' vector config."
                )
                self.sparse_enabled = False
        except Exception as exc:
            print(f"[WARN] Could not inspect collection sparse config for {COLLECTION_NAME}: {exc}")
            # Keep env-configured behavior if inspection fails.

    def _load_embedders(self) -> None:
        self._ensure_sparse_compatibility()
        if self._dense_embedder and (self._sparse_embedder or not self.sparse_enabled):
            return

        from fastembed import TextEmbedding  # type: ignore

        if not self._dense_embedder:
            print(f"[INIT] Loading Dense: {self.dense_model}")
            self._dense_embedder = TextEmbedding(model_name=self.dense_model)

        if self.sparse_enabled and not self._sparse_embedder:
            from fastembed import SparseTextEmbedding  # type: ignore
            print(f"[INIT] Loading Sparse: {self.sparse_model}")
            self._sparse_embedder = SparseTextEmbedding(model_name=self.sparse_model)

    def process_document(
        self,
        *args: str,
        file_path: str | None = None,
        tenant_id: str,
        domain: Domain = Domain.MATERIAL,
        document_id: str | None = None,
        visibility: str = "public",
        source_type: SourceType = SourceType.MANUAL,
        tags: list[str] | None = None,
    ) -> IngestStats:
        if args:
            if len(args) > 1:
                raise TypeError("process_document() takes at most 1 positional argument")
            if file_path is not None:
                raise TypeError("process_document() got multiple values for file_path")
            file_path = args[0]
        if not file_path:
            raise TypeError("process_document() missing required argument: file_path")
        _ensure_tenant_allowed(tenant_id)
        started = time.perf_counter()
        self._load_embedders()

        visibility = _coerce_visibility(visibility)
        filename = os.path.basename(file_path)
        doc_id = document_id or _hash_path(file_path)
        facets = _parse_facets_from_tags(tags)
        entity = facets.get("entity") or _guess_entity_from_filename(filename)
        aspects = facets.get("aspects") or []
        language = facets.get("language")
        source_version = facets.get("source_version")
        effective_date = facets.get("effective_date")
        material_code = _normalize_material_code(
            _extract_material_code(str(facets.get("material_code") or ""), filename)
        )
        source_url = str(facets.get("source_url") or file_path)
        shore_hardness_tag = facets.get("shore_hardness")
        try:
            shore_hardness_base = int(shore_hardness_tag) if shore_hardness_tag is not None else None
        except (TypeError, ValueError):
            shore_hardness_base = None
        if shore_hardness_base is None:
            shore_hardness_base = _extract_shore_hardness(filename)
        if shore_hardness_base is None:
            shore_hardness_base = _DEFAULT_SHORE_HARDNESS
        temp_range_payload = facets.get("temp_range")

        pages = _load_pages(file_path)
        if not pages or not any(text.strip() for _, text in pages):
            log.warning("ingest_skipped_empty_document", extra={"file_path": file_path, "file_name": filename})
            return IngestStats(chunks=0, elapsed_ms=int((time.perf_counter() - started) * 1000))

        document_text = "\n\n".join(text for _, text in pages if text and text.strip())
        material_code = _normalize_material_code(
            _extract_material_code(material_code, document_text, filename)
        )
        if shore_hardness_base is None:
            shore_hardness_base = _extract_shore_hardness(filename, document_text)
        try:
            shore_hardness_base = int(shore_hardness_base) if shore_hardness_base is not None else _DEFAULT_SHORE_HARDNESS
        except (TypeError, ValueError):
            shore_hardness_base = _DEFAULT_SHORE_HARDNESS
        temp_range_payload = _normalize_temp_range_payload(temp_range_payload, document_text)
        additional_metadata_doc = _extract_additional_metadata(
            text=document_text,
            filename=filename,
            tag_metadata=facets.get("additional_metadata") if isinstance(facets.get("additional_metadata"), dict) else None,
        )

        # V5.2 PED-Platinum: PDFs gehen AUSSCHLIESSLICH durch den strukturierten ETL-Pfad.
        # Kein Fallback auf rohes Chunking.
        is_platinum = filename.lower().endswith(".pdf")
        if is_platinum:
            log.info("platinum_etl_start", extra={"doc_filename": filename})
            # Hard-fail: Wenn LLM-Extraktion scheitert, schlägt der Ingest fehl.
            llm_extraction = _extract_platinum_structured_llm(text=document_text, filename=filename)
            etl_result = process_document_pipeline(llm_extraction, doc_id, additional_metadata_doc)

            if not etl_result.extracted_points:
                log.error(
                    "platinum_etl_no_points",
                    extra={
                        "doc_filename": filename,
                        "status": etl_result.status.value,
                        "quarantine_report": etl_result.quarantine_report,
                    },
                )
                raise ValueError(
                    f"Platinum ETL yielded 0 valid points for '{filename}'. "
                    f"Status: {etl_result.status.value}. "
                    f"Quarantine: {etl_result.quarantine_report}"
                )

            log.info(
                "platinum_etl_success",
                extra={"status": etl_result.status.value, "points": len(etl_result.extracted_points)},
            )
            # Baue raw_chunks aus ETL-Points (vector_text, page=None, etl_payload)
            raw_chunks = [(p["vector_text"], None, p) for p in etl_result.extracted_points]
        else:
            raw_chunks = _chunk_pages(pages, filename=filename)

        if not raw_chunks:
            log.warning("ingest_skipped_no_chunks", extra={"file_path": file_path, "file_name": filename})
            return IngestStats(chunks=0, elapsed_ms=int((time.perf_counter() - started) * 1000))

        dense_vecs = list(self._dense_embedder.embed([c[0] for c in raw_chunks]))  # type: ignore[union-attr]

        sparse_vecs = None
        if self.sparse_enabled:
            sparse_vecs = list(self._sparse_embedder.embed([c[0] for c in raw_chunks]))  # type: ignore[union-attr]

        points: List[models.PointStruct] = []
        bm25_docs: List[Document] = []
        seen_hashes: set[str] = set()
        for idx, chunk_tuple in enumerate(raw_chunks):
            # Platinum Metadata Overlay
            etl_payload = None
            if len(chunk_tuple) > 2:
                chunk_text, page_number, etl_payload = chunk_tuple
            else:
                chunk_text, page_number = chunk_tuple

            chunk_hash = ChunkMetadata.compute_hash(chunk_text)
            if chunk_hash in seen_hashes:
                continue
            seen_hashes.add(chunk_hash)
            chunk_id = ChunkMetadata.generate_chunk_id(tenant_id, doc_id, idx)

            chunk_additional_metadata = dict(additional_metadata_doc)

            if etl_payload:
                # Platinum-spezifische Metadaten injizieren
                chunk_additional_metadata.update({
                    "etl_status": etl_result.status.value,
                    "etl_quarantine_report": etl_result.quarantine_report,
                    "operating_point_data": etl_payload.get("limits", {}),
                    "operating_point_conditions": etl_payload.get("conditions", {}),
                    "pipeline_version": "V5.2-Platinum"
                })
                # Wir nutzen die normalisierten Werte aus ETL falls vorhanden
                chunk_shore = etl_payload.get("limits", {}).get("shore_hardness", {}).get("normalized", shore_hardness_base)
            else:
                chunk_shore = _extract_shore_hardness(chunk_text) or shore_hardness_base
                chunk_regex_metadata = _extract_dynamic_metadata_regex(text=chunk_text, filename=filename)
                for key, value in chunk_regex_metadata.items():
                    chunk_additional_metadata.setdefault(key, value)

            chunk_temp_range = _normalize_temp_range_payload(temp_range_payload, chunk_text)

            # Fix 2: eng.material_family aus additional_metadata befüllen
            _eng_material_family = str(chunk_additional_metadata.get("material_family") or "").strip() or None
            meta = ChunkMetadata(
                tenant_id=tenant_id,
                doc_id=doc_id,          # legacy
                document_id=doc_id,     # canonical
                chunk_id=chunk_id,
                chunk_hash=chunk_hash,
                source_uri=file_path,
                source_type=source_type,
                domain=domain,
                chunk_index=idx,
                entity=entity,
                aspect=aspects,
                language=language,
                source_version=source_version,
                effective_date=effective_date,
                material_code=material_code,
                source_url=source_url,
                shore_hardness=int(chunk_shore),
                temp_range=TempRange.model_validate(chunk_temp_range),
                additional_metadata=chunk_additional_metadata,
                title=filename,
                page_number=page_number,
                text=chunk_text,
                created_at=time.time(),
                visibility=visibility,
                eng=EngineeringProps(material_family=_eng_material_family),
            )
            meta_payload = meta.model_dump(mode="json")

            dense = dense_vecs[idx].tolist()

            # Vector shape must match the collection:
            # - unnamed collection => vector is list[float]
            # - named collection => vector is dict {name: list[float]} (plus optional sparse)
            if QDRANT_VECTOR_NAME:
                vector: Any
                if self.sparse_enabled and sparse_vecs is not None:
                    sparse = sparse_vecs[idx].as_object()
                    vector = {QDRANT_VECTOR_NAME: dense, "sparse": sparse}
                else:
                    vector = {QDRANT_VECTOR_NAME: dense}
            else:
                # Unnamed single-vector collection (like sealai-docs in your output)
                vector = dense

            # Platinum-PDFs erhalten document_meta für die State-Machine.
            # Standard-Chunks erhalten keinen document_meta (rückwärtskompatibel).
            point_payload: dict[str, Any] = {
                "metadata": meta_payload,
                "text": chunk_text,
                "source": str(file_path),
                "filename": filename,
                "tenant_id": tenant_id,
                "document_id": doc_id,
                "visibility": visibility,
                "material_code": meta_payload.get("material_code"),
                "source_url": meta_payload.get("source_url"),
                "shore_hardness": meta_payload.get("shore_hardness"),
                "temp_range": meta_payload.get("temp_range"),
                "additional_metadata": meta_payload.get("additional_metadata"),
            }
            if is_platinum:
                # Versioned status field für die bulletproof State-Machine
                point_payload["document_meta"] = {
                    "status": "VALIDATED",
                    "logical_document_key": doc_id,
                    "version_id": int(time.time()),
                    "pipeline_version": "V5.2-Platinum",
                    "etl_status": etl_result.status.value,
                    "quarantine_report": etl_result.quarantine_report,
                }
            points.append(
                models.PointStruct(
                    id=chunk_id,
                    vector=vector,
                    payload=point_payload,
                )
            )
            bm25_docs.append(
                Document(
                    page_content=chunk_text,
                    metadata={
                        **meta_payload,
                        "source": str(file_path),
                        "filename": filename,
                        "tenant_id": tenant_id,
                        "document_id": doc_id,
                        "chunk_id": chunk_id,
                    },
                    id=chunk_id,
                )
            )

        self.client.upsert(COLLECTION_NAME, points=points)

        # V5.2: Platinum-PDFs → atomisch auf PUBLISHED setzen und alte Versionen deprecaten
        if is_platinum and points:
            try:
                sm_result = transition_to_published_bulletproof(
                    client=self.client,
                    collection_name=COLLECTION_NAME,
                    logical_document_key=doc_id,
                    new_version_id=points[0].payload["document_meta"]["version_id"],
                    new_qdrant_id=str(points[0].id),
                )
                log.info(
                    "platinum_state_machine_success",
                    extra={
                        "operation_id": sm_result["operation_id"],
                        "published": sm_result["published_id"],
                        "deprecated_count": len(sm_result["deprecated_ids"]),
                    },
                )
            except Exception as sm_exc:
                # State-Machine-Fehler ist nicht kritisch: Daten sind bereits upsertet (VALIDATED).
                # HITL kann den Status manuell auf PUBLISHED setzen.
                log.warning(
                    "platinum_state_machine_failed",
                    extra={"doc_filename": filename, "error": str(sm_exc)},
                )
        try:
            from app.services.rag.bm25_store import bm25_repo

            bm25_repo.upsert_documents(COLLECTION_NAME, bm25_docs)
        except Exception as exc:
            print(f"[WARN] BM25 upsert skipped for {filename}: {type(exc).__name__}: {exc}")

        elapsed_ms = int((time.perf_counter() - started) * 1000)
        file_size = None
        try:
            file_size = os.path.getsize(file_path)
        except OSError:
            pass

        print(
            f"[INGEST] Upserted {len(points)} chunks -> {COLLECTION_NAME} "
            f"(tenant={tenant_id}, doc={doc_id}, vis={visibility}, vector_name={QDRANT_VECTOR_NAME!r}, sparse={self.sparse_enabled})"
        )
        return IngestStats(chunks=len(points), elapsed_ms=elapsed_ms, file_size=file_size)


# -----------------------------------------------------------------------------
# Worker-compatible API (IMPORTANT)
# -----------------------------------------------------------------------------

def ingest_file(
    file_path: str,
    *,
    tenant_id: str,
    document_id: str | None = None,
    category: str | None = None,
    tags: list[str] | None = None,
    visibility: str = "public",
    sha256: str | None = None,
    source: str | None = None,
    **kwargs: Any,
) -> Dict[str, Any]:
    """
    Backward-compatible entrypoint for the async job worker.
    """
    _ensure_tenant_allowed(tenant_id)
    domain = _parse_domain(category)
    canonical_doc_id = document_id or sha256 or _hash_path(file_path)
    filename = os.path.basename(file_path)
    facets = _parse_facets_from_tags(tags)
    material_code = _normalize_material_code(
        _extract_material_code(str(facets.get("material_code") or ""), filename)
    )
    source_url = str(facets.get("source_url") or source or file_path)
    shore_hardness_tag = facets.get("shore_hardness")
    try:
        shore_hardness_base = int(shore_hardness_tag) if shore_hardness_tag is not None else None
    except (TypeError, ValueError):
        shore_hardness_base = None
    if shore_hardness_base is None:
        shore_hardness_base = _extract_shore_hardness(filename)
    if shore_hardness_base is None:
        shore_hardness_base = _DEFAULT_SHORE_HARDNESS
    temp_range_payload = facets.get("temp_range")
    additional_metadata_seed = facets.get("additional_metadata") if isinstance(facets.get("additional_metadata"), dict) else {}

    # Legacy vectorstore path (optional): disabled by default to keep ingestion aligned
    # with retrieval embedding backend (fastembed).
    if LEGACY_VECTORSTORE_ENABLED and QdrantVectorStore is not None and HuggingFaceEmbeddings is not None:
        docs = load_document(file_path)
        if not docs:
            log.warning("ingest_skipped_empty_document", extra={"file_path": file_path, "file_name": filename})
            return {
                "ok": True,
                "file_path": file_path,
                "tenant_id": tenant_id,
                "document_id": canonical_doc_id,
                "category": category,
                "domain": getattr(domain, "value", str(domain)),
                "visibility": _coerce_visibility(visibility),
                "tags": tags or [],
                "sha256": sha256,
                "source": source,
                "chunks": 0,
                "elapsed_ms": 0,
                "file_size": None,
                "collection": COLLECTION_NAME,
                "dense_model": DENSE_MODEL,
                "sparse_model": SPARSE_MODEL if ENABLE_SPARSE else None,
                "vector_name": QDRANT_VECTOR_NAME,
                "sparse_enabled": ENABLE_SPARSE,
            }
        document_text = "\n\n".join(str(getattr(doc, "page_content", "") or "") for doc in docs)
        material_code = _normalize_material_code(
            _extract_material_code(material_code, filename, document_text)
        )
        shore_hardness_base = _extract_shore_hardness(filename, document_text) or shore_hardness_base or _DEFAULT_SHORE_HARDNESS
        temp_range_payload = _normalize_temp_range_payload(temp_range_payload, document_text)
        additional_metadata = _extract_additional_metadata(
            text=document_text,
            filename=filename,
            tag_metadata=additional_metadata_seed,
        )
        source_path = filename
        size_bytes = None
        try:
            size_bytes = os.path.getsize(file_path)
        except OSError:
            pass
        content_type = mimetypes.guess_type(file_path)[0] or "application/octet-stream"
        for doc in docs:
            meta = dict(getattr(doc, "metadata", {}) or {})
            section = meta.get("section") or meta.get("section_title")
            meta.update(
                {
                    "filename": filename,
                    "content_type": content_type,
                    "size_bytes": size_bytes,
                    "source_path": source_path,
                    "tenant_id": tenant_id,
                    "document_id": canonical_doc_id,
                    "visibility": _coerce_visibility(visibility),
                    "material_code": material_code,
                    "source_url": source_url,
                    "shore_hardness": _extract_shore_hardness(doc.page_content) or shore_hardness_base,
                    "temp_range": _normalize_temp_range_payload(temp_range_payload, doc.page_content),
                    "additional_metadata": additional_metadata,
                }
            )
            if section:
                meta["section"] = section
            doc.metadata = meta

        embeddings = HuggingFaceEmbeddings(model_name=DENSE_MODEL)
        QdrantVectorStore.from_documents(
            docs,
            embeddings,
            url=QDRANT_URL,
            api_key=QDRANT_API_KEY,
            collection_name=COLLECTION_NAME,
        )
        return {
            "ok": True,
            "file_path": file_path,
            "tenant_id": tenant_id,
            "document_id": canonical_doc_id,
            "category": category,
            "domain": getattr(domain, "value", str(domain)),
            "visibility": _coerce_visibility(visibility),
            "tags": tags or [],
            "sha256": sha256,
            "source": source,
            "chunks": len(docs),
            "elapsed_ms": 0,
            "file_size": size_bytes,
            "collection": COLLECTION_NAME,
            "dense_model": DENSE_MODEL,
            "sparse_model": SPARSE_MODEL if ENABLE_SPARSE else None,
            "vector_name": QDRANT_VECTOR_NAME,
            "sparse_enabled": ENABLE_SPARSE,
        }

    pipe = IngestPipeline()
    stats = pipe.process_document(
        file_path=file_path,
        tenant_id=tenant_id,
        domain=domain,
        document_id=canonical_doc_id,
        visibility=visibility,
        source_type=SourceType.MANUAL,
        tags=tags,
    )
    return {
        "ok": True,
        "file_path": file_path,
        "tenant_id": tenant_id,
        "document_id": canonical_doc_id,
        "category": category,
        "domain": getattr(domain, "value", str(domain)),
        "visibility": _coerce_visibility(visibility),
        "tags": tags or [],
        "sha256": sha256,
        "source": source,
        "chunks": stats.chunks,
        "elapsed_ms": stats.elapsed_ms,
        "file_size": stats.file_size,
        "collection": COLLECTION_NAME,
        "dense_model": DENSE_MODEL,
        "sparse_model": SPARSE_MODEL if pipe.sparse_enabled else None,
        "vector_name": QDRANT_VECTOR_NAME,
        "sparse_enabled": pipe.sparse_enabled,
    }


# -----------------------------------------------------------------------------
# CLI
# -----------------------------------------------------------------------------

def _iter_files(path: str) -> Iterable[str]:
    if os.path.isdir(path):
        for root, _, files in os.walk(path):
            for f in files:
                if f.lower().endswith((".docx", ".pdf", ".txt", ".md", ".log", ".csv")):
                    yield os.path.join(root, f)
    else:
        yield path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="SealAI RAG ingest into Qdrant.")
    parser.add_argument("path", help="File or Directory")
    parser.add_argument("--tenant", default=DEFAULT_INGEST_TENANT)
    parser.add_argument("--domain", default="material")
    parser.add_argument("--visibility", default="public", choices=["public", "private"])
    args = parser.parse_args()

    pipe = IngestPipeline()
    domain = _parse_domain(args.domain)
    visibility = _coerce_visibility(args.visibility)

    for p in _iter_files(args.path):
        pipe.process_document(
            file_path=p,
            tenant_id=args.tenant,
            domain=domain,
            document_id=None,
            visibility=visibility,
            source_type=SourceType.MANUAL,
        )
